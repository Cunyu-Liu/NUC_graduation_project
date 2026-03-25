#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中北大学软件学院本科毕业设计论文 DOCX 生成器
符合《中北大学本科生学位论文撰写格式》最新标准
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os


class NUCThesisConverter:
    """中北大学论文格式转换器"""
    
    def __init__(self):
        self.doc = Document()
        self.current_chapter = 0
        self.figure_counter = {}
        self.table_counter = {}
        self.in_abstract = False
        self.in_toc = False
        self.in_main = False
        
        self._setup_styles()
        
    def _setup_styles(self):
        """设置文档样式"""
        styles = self.doc.styles
        
        # 正文：小四号宋体(12pt)，1.5倍行距，首行缩进2字符(0.74cm)
        style_normal = styles['Normal']
        style_normal.font.name = 'Times New Roman'
        style_normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style_normal.font.size = Pt(12)
        style_normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        style_normal.paragraph_format.first_line_indent = Cm(0.74)
        style_normal.paragraph_format.space_before = Pt(0)
        style_normal.paragraph_format.space_after = Pt(0)
        
        # 一级标题：小三号黑体(15pt)，加粗，段前段后0.5行(6pt)
        style_h1 = styles.add_style('CustomH1', WD_STYLE_TYPE.PARAGRAPH)
        style_h1.font.name = 'Times New Roman'
        style_h1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style_h1.font.size = Pt(15)
        style_h1.font.bold = True
        style_h1.paragraph_format.space_before = Pt(6)
        style_h1.paragraph_format.space_after = Pt(6)
        style_h1.paragraph_format.first_line_indent = Cm(0)
        style_h1.paragraph_format.keep_with_next = True
        
        # 二级标题：小四号黑体(12pt)，加粗
        style_h2 = styles.add_style('CustomH2', WD_STYLE_TYPE.PARAGRAPH)
        style_h2.font.name = 'Times New Roman'
        style_h2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style_h2.font.size = Pt(12)
        style_h2.font.bold = True
        style_h2.paragraph_format.space_before = Pt(6)
        style_h2.paragraph_format.space_after = Pt(6)
        style_h2.paragraph_format.first_line_indent = Cm(0)
        style_h2.paragraph_format.keep_with_next = True
        
        # 三级标题：小四号黑体(12pt)，不加粗
        style_h3 = styles.add_style('CustomH3', WD_STYLE_TYPE.PARAGRAPH)
        style_h3.font.name = 'Times New Roman'
        style_h3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style_h3.font.size = Pt(12)
        style_h3.font.bold = False
        style_h3.paragraph_format.space_before = Pt(3)
        style_h3.paragraph_format.space_after = Pt(3)
        style_h3.paragraph_format.first_line_indent = Cm(0)
        style_h3.paragraph_format.keep_with_next = True
        
        # 四级标题：小四号黑体(12pt)，不加粗，缩进
        style_h4 = styles.add_style('CustomH4', WD_STYLE_TYPE.PARAGRAPH)
        style_h4.font.name = 'Times New Roman'
        style_h4._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style_h4.font.size = Pt(12)
        style_h4.font.bold = False
        style_h4.paragraph_format.first_line_indent = Cm(0.74)
        
        # 图表标题：五号宋体(10.5pt)，居中
        try:
            style_cap = styles.add_style('CustomCaption', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style_cap = styles['CustomCaption']
        style_cap.font.name = 'Times New Roman'
        style_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style_cap.font.size = Pt(10.5)
        style_cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_cap.paragraph_format.space_before = Pt(3)
        style_cap.paragraph_format.space_after = Pt(3)
        
        # 代码块：五号(10.5pt)，单倍行距，灰色底纹
        try:
            style_code = styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style_code = styles['CodeBlock']
        style_code.font.name = 'Courier New'
        style_code._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style_code.font.size = Pt(10.5)
        style_code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style_code.paragraph_format.left_indent = Cm(0.5)
        style_code.paragraph_format.first_line_indent = Cm(0)
        
        # 摘要标题：小三号黑体，居中
        try:
            style_abs_title = styles.add_style('AbstractTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style_abs_title = styles['AbstractTitle']
        style_abs_title.font.name = 'Times New Roman'
        style_abs_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style_abs_title.font.size = Pt(15)
        style_abs_title.font.bold = True
        style_abs_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_abs_title.paragraph_format.space_after = Pt(12)
        
        # 目录标题：四号黑体(14pt)，居中
        try:
            style_toc_title = styles.add_style('TocTitle', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            style_toc_title = styles['TocTitle']
        style_toc_title.font.name = 'Times New Roman'
        style_toc_title._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style_toc_title.font.size = Pt(14)
        style_toc_title.font.bold = True
        style_toc_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_toc_title.paragraph_format.space_after = Pt(12)
        
    def _setup_section(self, section):
        """设置节格式"""
        # A4纸张：21cm x 29.7cm
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        # 页边距：上30mm，下25mm，左30mm，右20mm
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        # 页眉页脚距离
        section.header_distance = Cm(2.5)
        section.footer_distance = Cm(1.8)
        
    def _add_section_break(self, start_type=WD_SECTION_START.NEW_PAGE):
        """插入分节符"""
        self.doc.add_section(start_type)
        section = self.doc.sections[-1]
        self._setup_section(section)
        return section
        
    def _set_no_page_number(self, section):
        """设置该节无页码"""
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), '0')
        sectPr.append(pgNumType)
        
    def _set_page_number_arabic(self, section, start=1):
        """设置该节使用阿拉伯数字页码"""
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), str(start))
        pgNumType.set(qn('w:fmt'), 'decimal')
        sectPr.append(pgNumType)
        
        # 添加页脚页码
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = ""
        
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(10.5)
        
    def _set_page_number_roman(self, section, start=1):
        """设置该节使用罗马数字页码"""
        sectPr = section._sectPr
        pgNumType = OxmlElement('w:pgNumType')
        pgNumType.set(qn('w:start'), str(start))
        pgNumType.set(qn('w:fmt'), 'upperRoman')
        sectPr.append(pgNumType)
        
        # 添加页脚页码
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = ""
        
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(10.5)
        
    def _add_header(self, section, text="中北大学2026届毕业设计说明书"):
        """添加页眉"""
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.text = text
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_para.runs[0] if header_para.runs else header_para.add_run(text)
        run.font.size = Pt(12)  # 小四号
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 添加下划线
        pPr = header_para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        
    def convert(self, md_path, out_path):
        """转换Markdown到DOCX"""
        if not os.path.exists(md_path):
            raise FileNotFoundError(f"文件不存在: {md_path}")
            
        with open(md_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        # 初始化第一节
        first_section = self.doc.sections[0]
        self._setup_section(first_section)
        self._set_no_page_number(first_section)
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 识别中文摘要
            if line in ['中文摘要', '## 中文摘要', '# 中文摘要']:
                sect = self._add_section_break()
                self._set_no_page_number(sect)
                self._add_header(sect)
                self._add_abstract_title('中文摘要')
                self.in_abstract = True
                self.in_toc = False
                self.in_main = False
                i += 1
                continue
            
            # 识别英文摘要
            if line in ['Abstract', '## Abstract', '# Abstract']:
                sect = self._add_section_break()
                self._set_no_page_number(sect)
                self._add_header(sect)
                self._add_abstract_title('Abstract')
                i += 1
                continue
            
            # 识别目录
            if line in ['目录', '## 目录', '# 目录']:
                sect = self._add_section_break()
                self._set_page_number_roman(sect, 1)  # 目录使用罗马数字
                self._add_header(sect)
                self._add_toc_title('目录')
                self.in_toc = True
                self.in_abstract = False
                self.in_main = False
                i += 1
                continue
            
            # 识别主要符号表
            if line in ['主要符号表', '## 主要符号表']:
                sect = self._add_section_break()
                if not self.in_main:
                    self._set_page_number_roman(sect)  # 继续使用罗马数字
                self._add_header(sect)
                self._add_heading('主要符号表', 1)
                i += 1
                continue
            
            # 识别第1章/引言（正文开始）
            if self._is_chapter_start(line):
                if not self.in_main:
                    # 正文新节，使用阿拉伯数字页码
                    sect = self._add_section_break()
                    self._set_page_number_arabic(sect, 1)
                    self._add_header(sect)
                    self.in_main = True
                    self.in_toc = False
                    self.in_abstract = False
                
                self.current_chapter += 1
                self.figure_counter[self.current_chapter] = 0
                self.table_counter[self.current_chapter] = 0
                
                # 添加章标题
                p = self.doc.add_paragraph()
                p.style = 'CustomH1'
                run = p.add_run(line.replace('# ', ''))
                
                i += 1
                continue
            
            # 识别参考文献
            if line in ['参考文献', '## 参考文献']:
                sect = self._add_section_break()
                self._add_header(sect)
                self._add_heading('参考文献', 1)
                i += 1
                continue
            
            # 识别致谢
            if line in ['致谢', '## 致谢']:
                sect = self._add_section_break()
                self._add_header(sect)
                self._add_heading('致谢', 1)
                i += 1
                continue
            
            # 识别附录
            if line.startswith('附录') or line.startswith('## 附录'):
                if not self.in_main:
                    sect = self._add_section_break()
                    self._add_header(sect)
                self._add_heading(line.replace('## ', ''), 1)
                i += 1
                continue
            
            # 处理一般一级标题（# 标题）
            if self._is_level1_heading(line):
                p = self.doc.add_paragraph()
                p.style = 'CustomH1'
                run = p.add_run(line[2:])  # 去掉 '# '
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(15)  # 小三号
                run.font.bold = True
                i += 1
                continue
            
            # 处理二级标题
            if line.startswith('## '):
                p = self.doc.add_paragraph()
                p.style = 'CustomH2'
                run = p.add_run(line[3:])
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(12)  # 小四号
                run.font.bold = True
                i += 1
                continue
            
            # 处理三级标题
            if line.startswith('### '):
                p = self.doc.add_paragraph()
                p.style = 'CustomH3'
                run = p.add_run(line[4:])
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(12)  # 小四号
                run.font.bold = False
                i += 1
                continue
            
            # 处理四级标题
            if line.startswith('#### '):
                p = self.doc.add_paragraph()
                p.style = 'CustomH4'
                run = p.add_run(line[5:])
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(12)  # 小四号
                run.font.bold = False
                i += 1
                continue
            
            # 处理无序列表项 (- 或 * 开头) - 转换为普通段落，保持自然段格式
            if re.match(r'^[-\*]\s+', line):
                content = re.sub(r'^[-\*]\s+', '', line).strip()
                # 检查是否是连续的多行列表项，合并为一个段落
                list_contents = [content]
                j = i + 1
                while j < len(lines) and re.match(r'^[-\*]\s+', lines[j]):
                    list_contents.append(re.sub(r'^[-\*]\s+', '', lines[j]).strip())
                    j += 1
                
                # 合并为一个自然段，用分号连接（在同一run中，不创建新段落）
                combined_text = '；'.join(list_contents) + '。'
                p = self.doc.add_paragraph()
                p.style = 'Normal'
                # 直接使用add_run，不调用_add_formatted_text以避免复杂的分段逻辑
                run = p.add_run(combined_text)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
                i = j
                continue
            
            # 处理有序列表项 (1. 2. 等开头) - 转换为普通段落
            if re.match(r'^\d+\.\s+', line):
                content = re.sub(r'^\d+\.\s+', '', line).strip()
                # 检查是否是连续的多行有序列表，合并为一个段落
                list_contents = [content]
                j = i + 1
                while j < len(lines) and re.match(r'^\d+\.\s+', lines[j]):
                    list_contents.append(re.sub(r'^\d+\.\s+', '', lines[j]).strip())
                    j += 1
                
                # 合并为一个自然段，用分号连接
                combined_text = '；'.join(list_contents) + '。'
                p = self.doc.add_paragraph()
                p.style = 'Normal'
                run = p.add_run(combined_text)
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
                i = j
                continue
            
            # 处理代码块
            if line.startswith('```'):
                code_lines = []
                lang = line[3:].strip()
                i += 1
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                self._add_code('\n'.join(code_lines), lang)
                i += 1
                continue
            
            # 处理表格
            if line.startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_table(table_lines)
                continue
            
            # 处理图片
            if line.startswith('!['):
                self._add_image(line)
                i += 1
                continue
            
            # 处理图表标题（单独出现的图表标题）
            if re.match(r'^(图|表)\s*[\dX]', line) or line.startswith('**图') or line.startswith('**表'):
                # 检查下一行是否是表格或图片，如果不是则作为独立标题处理
                next_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
                if not next_line.startswith('|') and not next_line.startswith('!['):
                    self._add_caption(line.replace('**', ''))
                # 如果下一行是表格，标题已在表格处理中处理，跳过
                i += 1
                continue
            
            # 处理关键词
            if line.startswith('**关键词：**') or line.startswith('**Keywords：**'):
                p = self.doc.add_paragraph()
                p.style = 'Normal'
                p.paragraph_format.first_line_indent = Cm(0)
                run = p.add_run(line.replace('**', ''))
                run.bold = True
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                i += 1
                continue
            
            # 处理正文
            if line:
                self._add_text(line)
            
            i += 1
        
        self.doc.save(out_path)
        print(f"✅ 已生成: {out_path}")
        print("⚠️  请检查：1.页码是否正确 2.三线表样式 3.图片位置")
        
    def _is_chapter_start(self, line):
        """识别章标题"""
        patterns = [
            r'^第[一二三四五六七八九十\d]+章',
            r'^第\s*\d+\s*章',
            r'^#\s*第[一二三四五六七八九十\d]+章',
            r'^#\s*引言',
            r'^#\s*绪论',
            r'^#\s*结论',
        ]
        return any(re.match(p, line) for p in patterns)
    
    def _is_level1_heading(self, line):
        """识别一级标题（# 开头的非特殊标题）"""
        # 匹配 # 标题 但不是 ## 开头（二级标题）
        if line.startswith('# ') and not line.startswith('## '):
            # 排除已经处理过的特殊标题
            special_titles = ['引言', '绪论', '结论', '第']
            title_content = line[2:].strip()
            return not any(title_content.startswith(st) for st in special_titles)
        return False
        
    def _add_abstract_title(self, text):
        """添加摘要标题"""
        p = self.doc.add_paragraph()
        p.style = 'AbstractTitle'
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
    def _add_toc_title(self, text):
        """添加目录标题"""
        p = self.doc.add_paragraph()
        p.style = 'TocTitle'
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
    def _add_heading(self, text, level):
        """添加标题"""
        if level == 1:
            style = 'CustomH1'
        elif level == 2:
            style = 'CustomH2'
        elif level == 3:
            style = 'CustomH3'
        else:
            style = 'CustomH4'
        p = self.doc.add_paragraph()
        p.style = style
        p.add_run(text)
        
    def _add_text(self, text):
        """添加正文"""
        p = self.doc.add_paragraph()
        p.style = 'Normal'
        self._add_formatted_text(p, text)
    
    def _add_formatted_text(self, para, text):
        """添加带格式的文本 - 简化版，优先保证正确性"""
        import re
        
        # 简单方法：逐个字符扫描，识别 ** 和 * 标记
        i = 0
        n = len(text)
        current_text = ""
        is_bold = False
        is_italic = False
        
        while i < n:
            # 检查 ** (粗体)
            if i + 1 < n and text[i:i+2] == '**':
                # 先添加当前累积的文本
                if current_text:
                    run = para.add_run(current_text)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(12)
                    if is_bold:
                        run.bold = True
                    if is_italic:
                        run.italic = True
                    current_text = ""
                # 切换粗体状态
                is_bold = not is_bold
                i += 2
                continue
            
            # 检查 * (斜体) - 确保不是 ** 的一部分
            if text[i] == '*' and (i == 0 or text[i-1] != '*') and (i + 1 >= n or text[i+1] != '*'):
                # 先添加当前累积的文本
                if current_text:
                    run = para.add_run(current_text)
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(12)
                    if is_bold:
                        run.bold = True
                    if is_italic:
                        run.italic = True
                    current_text = ""
                # 切换斜体状态
                is_italic = not is_italic
                i += 1
                continue
            
            # 普通字符
            current_text += text[i]
            i += 1
        
        # 添加剩余的文本
        if current_text:
            run = para.add_run(current_text)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)
            if is_bold:
                run.bold = True
            if is_italic:
                run.italic = True
                
    def _set_three_line_table_style(self, table):
        """设置三线表样式：顶线、表头底线、底线"""
        from docx.oxml import parse_xml
        from docx.oxml.ns import qn as oxml_qn
        
        # 获取表格的XML元素
        tbl = table._tbl
        
        # 设置表格边框为无
        tblPr = tbl.tblPr
        
        # 移除已有的边框设置
        for child in list(tblPr):
            if child.tag.endswith('tblBorders'):
                tblPr.remove(child)
        
        # 创建三线表边框设置
        # 顶线和底线使用1.5磅，表头底线使用0.75磅
        borders_xml = parse_xml(r'''
            <w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>
                <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>
                <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            </w:tblBorders>
        ''')
        
        tblPr.append(borders_xml)
        
        # 为表头行（第一行）添加底线
        if len(table.rows) > 0:
            first_row = table.rows[0]
            row_borders = parse_xml(r'''
                <w:trPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:trBorders>
                        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="000000"/>
                    </w:trBorders>
                </w:trPr>
            ''')
            first_row._tr.get_or_add_trPr().append(row_borders)
    
    def _add_code(self, code, lang):
        """添加代码块 - 作为普通正文处理，不加灰色底纹"""
        if not code.strip():
            return
            
        # 将代码作为普通段落添加，不使用代码块样式
        p = self.doc.add_paragraph()
        p.style = 'Normal'
        
        # 添加代码文本，使用等宽字体但不加灰色底纹
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(10.5)
        
    def _add_table(self, lines):
        """添加三线表 - 符合中北大学毕业设计格式要求"""
        if len(lines) < 2:
            return
            
        # 解析表格
        rows = []
        for line in lines:
            cells = [c.strip() for c in line.strip('|').split('|')]
            rows.append(cells)
        
        # 跳过分隔行（Markdown表格中的分隔线如 |---|---|）
        data_rows = [rows[0]] + [r for r in rows[1:] if not re.match(r'^[\s\-\|]+$', '|'.join(r))]
        
        if not data_rows or len(data_rows) < 1:
            return
            
        num_cols = len(data_rows[0])
        if num_cols == 0:
            return
        
        # 提取表格标题（第一行第一列如果包含"表X.X"或作为表名）
        table_title = ""
        header_row_idx = 0
        
        # 检查第一行是否是表标题行
        first_cell = data_rows[0][0] if data_rows[0] else ""
        if re.match(r'^表[\dX][.．][\dX]', first_cell) or len(data_rows) == 1:
            # 第一行包含表号，作为表标题
            table_title = first_cell
            data_rows = data_rows[1:]  # 移除标题行，保留实际数据行
            header_row_idx = 0
        
        if not data_rows:
            return
            
        # 如果有标题，添加表标题（表上方）
        if table_title:
            cap_p = self.doc.add_paragraph()
            cap_p.style = 'CustomCaption'
            cap_p.paragraph_format.space_after = Pt(6)
            cap_p.paragraph_format.space_before = Pt(6)
            run = cap_p.add_run(table_title)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(10.5)  # 五号
        
        # 创建表格
        table = self.doc.add_table(rows=len(data_rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 填充数据
        for i, row_data in enumerate(data_rows):
            row = table.rows[i]
            for j, text in enumerate(row_data[:num_cols]):
                cell = row.cells[j]
                cell.text = text
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                
                # 设置字体 - 五号宋体
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(10.5)  # 五号
                        # 第一行（表头）加粗
                        if i == 0:
                            run.font.bold = True
        
        # 设置三线表样式
        self._set_three_line_table_style(table)
        
    def _add_image(self, line):
        """处理图片"""
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if not match:
            return
            
        alt, path = match.groups()
        
        # 插入图片
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if os.path.exists(path):
            try:
                run = p.add_run()
                run.add_picture(path, width=Inches(5.0))
            except:
                p.add_run(f"[图片: {alt}]").font.color.rgb = RGBColor(255,0,0)
        else:
            p.add_run(f"[图片占位: {alt}]").font.color.rgb = RGBColor(255,0,0)
        
        # 图注（下方）
        self.figure_counter[self.current_chapter] = self.figure_counter.get(self.current_chapter, 0) + 1
        fig_num = self.figure_counter[self.current_chapter]
        
        cap_p = self.doc.add_paragraph()
        cap_p.style = 'CustomCaption'
        cap_p.paragraph_format.space_before = Pt(6)
        cap_p.add_run(f"图{self.current_chapter}.{fig_num} {alt}")
        
    def _add_caption(self, text):
        """添加图表标题"""
        p = self.doc.add_paragraph()
        p.style = 'CustomCaption'
        p.add_run(text)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) >= 2:
        md_file = sys.argv[1]
        docx_file = sys.argv[2] if len(sys.argv) >= 3 else md_file.replace('.md', '.docx')
        
        converter = NUCThesisConverter()
        converter.convert(md_file, docx_file)
    else:
        print("用法: python md-docx.py input.md [output.docx]")
        print("示例: python md-docx.py 毕业设计说明书_精简版.md 毕业设计说明书.docx")
