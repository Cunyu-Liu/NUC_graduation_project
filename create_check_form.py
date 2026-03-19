#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
创建毕业设计中期检查表
"""
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'single')
            edge_element.set(qn('w:sz'), '4')
            edge_element.set(qn('w:space'), '0')
            edge_element.set(qn('w:color'), '000000')
            tcPr.append(edge_element)

def set_font(run, font_name='宋体', font_size=10.5, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def create_check_form():
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("中北大学软件学院毕业设计中期检查表")
    set_font(run, font_name='黑体', font_size=16, bold=True)
    
    # 添加空行
    doc.add_paragraph()
    
    # 创建基本信息表格
    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'
    
    # 设置表格内容
    # 第一行
    row = info_table.rows[0]
    row.cells[0].text = "专业"
    row.cells[1].text = "软件工程"
    row.cells[2].text = "方向"
    row.cells[3].text = "物联网与智慧城市建设"
    
    # 第二行
    row = info_table.rows[1]
    row.cells[0].text = "班级"
    row.cells[1].text = "22130415"
    row.cells[2].text = "姓名"
    row.cells[3].text = "刘存宇"
    
    # 第三行
    row = info_table.rows[2]
    row.cells[0].text = "学号"
    row.cells[1].text = "2213041523"
    row.cells[2].text = "指导教师"
    row.cells[3].text = "王斌"
    
    # 第四行
    row = info_table.rows[3]
    row.cells[0].text = "校外指导教师姓名"
    row.cells[1].text = ""
    row.cells[2].text = "题目"
    row.cells[3].text = "科研文献摘要提取系统的设计与实现"
    
    # 设置所有单元格字体
    for row in info_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font(run, font_name='宋体', font_size=10.5)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 设置列宽
    for row in info_table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(6)
    
    doc.add_paragraph()
    
    # 添加"本人在该设计中具体应完成的工作"部分
    section1_title = doc.add_paragraph()
    run = section1_title.add_run("本人在该设计中具体应完成的工作")
    set_font(run, font_name='黑体', font_size=12, bold=True)
    
    section1_content = """1. 完成科研文献摘要提取系统的需求分析与系统设计
2. 实现PDF文档的智能解析与文本提取功能
3. 基于GLM-4大语言模型实现文献摘要自动生成
4. 实现12类核心要点（创新点、方法、实验、结论等）的智能提取
5. 构建知识图谱可视化模块，展示论文间关联关系
6. 实现基于研究空白的智能代码生成功能
7. 完成系统的前后端开发与联调
8. 撰写毕业设计说明书"""
    
    section1_para = doc.add_paragraph()
    run = section1_para.add_run(section1_content)
    set_font(run, font_name='宋体', font_size=10.5)
    section1_para.paragraph_format.line_spacing = 1.5
    section1_para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    
    # 添加"简述毕业设计开始以来所做的具体工作和取得的进展"部分
    section2_title = doc.add_paragraph()
    run = section2_title.add_run("简述毕业设计开始以来所做的具体工作和取得的进展")
    set_font(run, font_name='黑体', font_size=12, bold=True)
    
    section2_content = """一、已完成的核心功能开发
1. PDF智能解析模块：实现对中英文PDF文档的高精度解析，支持文本、表格、图片提取
2. AI摘要生成模块：基于GLM-4 API实现博士级学术摘要自动生成
3. 要点提取模块：智能提取论文中的12类核心要点，包括创新点、研究方法、实验设计、关键结论等
4. 研究空白挖掘模块：自动识别5种类型的研究空白，并进行优先级排序

二、知识图谱与可视化
1. 基于D3.js实现力导向布局的知识图谱可视化
2. 自动构建论文关系网络，支持5种关系类型的识别与展示
3. 实现节点交互、关系筛选、缩放拖拽等功能

三、AI智能助手功能
1. Kimi风格AI聊天系统：实现流式输出、Markdown渲染、LaTeX公式支持
2. 基于RAG技术的论文库智能问答
3. LangChain链式工作流：支持SequentialChain多步骤分析流程

四、向量聚类与推荐
1. 集成Milvus向量数据库，实现基于BGE-large模型的语义嵌入
2. 实现论文语义相似度聚类分析
3. 支持相似论文智能推荐

五、代码生成引擎
1. 实现6种代码生成策略：方法改进、新方法提出、数据集创建、实验设计、模型实现、算法优化
2. 集成Monaco Editor代码编辑器
3. 支持代码版本历史管理

六、性能优化
1. 异步工作流引擎：支持100篇论文并发处理，分析速度提升6倍
2. 数据库优化：创建30+索引，查询速度提升10-50倍
3. Redis缓存层：提升系统响应速度"""
    
    section2_para = doc.add_paragraph()
    run = section2_para.add_run(section2_content)
    set_font(run, font_name='宋体', font_size=10.5)
    section2_para.paragraph_format.line_spacing = 1.5
    section2_para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    
    # 添加"目前存在问题，下一步的主要设计任务，具体设想与安排"部分
    section3_title = doc.add_paragraph()
    run = section3_title.add_run("目前存在问题，下一步的主要设计任务，具体设想与安排")
    set_font(run, font_name='黑体', font_size=12, bold=True)
    
    section3_content = """一、目前存在的问题
1. 系统测试覆盖度有待提升，部分边界情况处理不够完善
2. 大模型API调用成本较高，需要优化调用策略
3. 知识图谱在大规模数据下的渲染性能有待优化
4. 部分复杂PDF格式的解析准确率需要进一步提升

二、下一步主要设计任务
1. 系统测试与优化（3月中旬完成）
   - 完善单元测试和集成测试
   - 进行压力测试和性能调优
   - 修复测试过程中发现的Bug

2. 论文撰写（3月下旬至4月中旬）
   - 完成毕业设计说明书初稿
   - 完善系统使用说明文档
   - 整理项目开发文档

3. 系统部署（4月中旬完成）
   - 完成Docker容器化部署配置
   - 编写部署文档和用户手册
   - 准备演示环境

4. 答辩准备（4月下旬完成）
   - 制作答辩PPT
   - 准备系统演示
   - 进行答辩演练

5. 最终验收（5月上旬）
   - 提交所有材料
   - 参加毕业答辩
   - 根据反馈进行修改完善"""
    
    section3_para = doc.add_paragraph()
    run = section3_para.add_run(section3_content)
    set_font(run, font_name='宋体', font_size=10.5)
    section3_para.paragraph_format.line_spacing = 1.5
    section3_para.paragraph_format.space_after = Pt(12)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加指导教师评价部分
    section4_title = doc.add_paragraph()
    run = section4_title.add_run("指导教师对该学生前期设计工作的评价")
    set_font(run, font_name='黑体', font_size=12, bold=True)
    
    section4_note = doc.add_paragraph()
    run = section4_note.add_run("（指导教师填写，必须有具体的评价意见，然后写'是否同意继续设计工作'）")
    set_font(run, font_name='宋体', font_size=10)
    run.font.italic = True
    
    # 添加签字行
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    sign_para = doc.add_paragraph()
    sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = sign_para.add_run("指导教师签字：                        年    月    日")
    set_font(run, font_name='宋体', font_size=10.5)
    
    # 保存文档
    output_path = "/Users/liucunyu/Documents/all_code/NUC_graduation_project/毕业设计说明书_完整版/2213041523-刘存宇-毕业设计中期检查表.docx"
    doc.save(output_path)
    print(f"中期检查表已保存至: {output_path}")

if __name__ == "__main__":
    create_check_form()
